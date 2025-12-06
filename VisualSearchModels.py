import numpy as np
import pandas as pd
import random
from scipy.optimize import minimize
from scipy.stats import chi2, multivariate_t
from concurrent.futures import ProcessPoolExecutor
from utils.DualProcess import generate_random_trial_sequence
import time
import ast
from scipy.special import psi
from scipy.stats import dirichlet
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX


MODEL_BOUNDS = {
    # 2-parameter RL models
    'decay': [(0.0001, 4.9999), (0.0001, 0.9999)],
    'delta': [(0.0001, 4.9999), (0.0001, 0.9999)],
    'decay_choice': [(0.0001, 4.9999), (0.0001, 0.9999)],
    'decay_win': [(0.0001, 4.9999), (0.0001, 0.9999)],
    'delta_RPUT': [(0.0001, 4.9999), (0.0001, 0.9999)],
    'decay_RPUT': [(0.0001, 4.9999), (0.0001, 0.9999)],
    'delta_RPUT_unc': [(0.0001, 4.9999), (0.0001, 0.9999), (-10, 10)],
    'decay_RPUT_unc': [(0.0001, 4.9999), (0.0001, 0.9999), (-10, 10)],
    'delta_perseveration': [(0.0001, 4.9999), (0.0001, 0.9999), (0.0001, 9.9999)],
    'delta_uncertainty': [(0.0001, 4.9999), (0.0001, 0.9999), (-10, 10)],
    'decay_uncertainty': [(0.0001, 4.9999), (0.0001, 0.9999), (-10, 10)],

    # PVL family
    'delta_PVL': [(0.0001, 4.9999), (0.0001, 0.9999), (0.0001, 0.9999), (0.0001, 4.9999)],
    'delta_PVL_relative': [(0.0001, 4.9999), (0.0001, 0.9999), (0.0001, 0.9999), (0.0001, 4.9999)],
    'decay_PVL_relative': [(0.0001, 4.9999), (0.0001, 0.9999), (0.0001, 0.9999), (0.0001, 4.9999)],

    # asymmetric delta
    'delta_asymmetric': [(0.0001, 4.9999), (0.0001, 0.9999), (0.0001, 0.9999)],

    # mean–variance utility model
    'mean_var': [(0.0001, 4.9999), (-100, 100)],
    'mean_var_delta': [(0.0001, 4.9999), (0.0001, 0.9999), (-10, 10)],
    'mean_var_decay': [(0.0001, 4.9999), (0.0001, 0.9999), (-10, 10)],
    'mean_var_unc': [(0.0001, 4.9999), (0.0001, 0.9999), (-10, 10)],
    'kalman_filter': [(0.0001, 4.9999), (0.0001, 0.9999)],

    # ACT-R
    'ACTR': [(0.0001, 4.9999), (0.0001, 0.9999), (-1.9999, -0.0001)],
    'ACTR_Ori': [(0.0001, 0.9999), (0.0001, 0.9999), (-1.9999, -0.0001)],

    # WSLS
    'perseveration': [(0.0001, 0.9999)],
    'WSLS': [(0.0001, 0.9999), (0.0001, 0.9999)],
    'WSLS_delta': [(0.0001, 0.9999), (0.0001, 0.9999), (0.0001, 0.9999)],
    'WSLS_delta_weight': [(0.0001, 4.9999), (0.0001, 0.9999),
                          (0.0001, 0.9999), (0.0001, 0.9999), (0.0001, 0.9999)],
    'WSLS_decay_weight': [(0.0001, 4.9999), (0.0001, 0.9999),
                          (0.0001, 0.9999), (0.0001, 0.9999), (0.0001, 0.9999)],

    # RT models
    'RT_exp_basic': [(0.0001, 4.9999), (0.0, 0.01), (0.3000, 23.9999)],
    'RT_delta': [(0.0001, 4.9999), (0.0001, 0.9999), (0.3000, 23.9999)],
    'RT_decay': [(0.0001, 4.9999), (0.0001, 4.9999), (0.3000, 23.9999)],
    'RT_exp_delta': [(0.0001, 4.9999), (0.0001, 0.9999),
                     (0.3000, 23.9999), (0.0, 0.01)],
    'RT_delta_PVL': [(0.0001, 4.9999), (0.0001, 0.9999),
                     (0.3000, 23.9999), (0.0001, 0.9999), (0.0001, 4.9999)],
    'RT_decay_PVL': [(0.0001, 4.9999), (0.0001, 4.9999),
                     (0.3000, 23.9999), (0.0001, 0.9999), (0.0001, 4.9999)],

    # hybrid models
    'hybrid_delta_delta': [(0.0001, 4.9999), (0.0001, 0.9999),
                           (0.3000, 23.9999), (0.0001, 0.9999)],
    'hybrid_decay_delta': [(0.0001, 4.9999), (0.0001, 0.9999),
                           (0.3000, 23.9999), (0.0001, 0.9999)],
    'hybrid_decay_decay': [(0.0001, 4.9999), (0.0001, 0.9999),
                           (0.3000, 23.9999), (0.0001, 0.9999)],

    'hybrid_delta_delta_3': [(0.0001, 4.9999), (0.0001, 0.9999),
                             (0.0001, 0.9999), (0.3000, 23.9999), (0.0001, 0.9999)],
    'hybrid_decay_delta_3': [(0.0001, 4.9999), (0.0001, 0.9999),
                             (0.0001, 0.9999), (0.3000, 23.9999), (0.0001, 0.9999)],
    'hybrid_decay_decay_3': [(0.0001, 4.9999), (0.0001, 0.9999),
                             (0.0001, 0.9999), (0.3000, 23.9999), (0.0001, 0.9999)],
}


def random_initial_guess(bounds):
    return [np.random.uniform(low, high) for (low, high) in bounds]


def fit_participant(model, participant_id, pdata, model_type, num_iterations=100,
                    beta_lower=-1, beta_upper=1):
    print(f"Fitting participant {participant_id}...")
    start_time = time.time()

    total_n = len(pdata['reward'])  # Total number of trials

    model.iteration = 0
    best_nll = 100000  # Initialize best negative log likelihood to a large number
    best_initial_guess = None
    best_parameters = None
    best_EV = None
    best_RT = None

    for _ in range(num_iterations):  # Randomly initiate the starting parameter for 1000 times

        model.iteration += 1

        print('Participant {} - Iteration [{}/{}]'.format(participant_id, model.iteration,
                                                          num_iterations))

        bounds = MODEL_BOUNDS[model_type]

        # generate initial guesses from dictionary bounds
        initial_guess = random_initial_guess(bounds)

        result = minimize(model.negative_log_likelihood, initial_guess,
                          args=(pdata['reward'], pdata['choice'], pdata['react_time']),
                          bounds=bounds, method='L-BFGS-B', options={'maxiter': 10000})

        if result.fun < best_nll:
            best_nll = result.fun
            best_initial_guess = initial_guess
            best_parameters = result.x
            best_EV = model.EVs.copy()
            best_RT = model.RTs.copy()

    k = len(best_parameters)  # Number of parameters
    aic = 2 * k + 2 * best_nll
    bic = k * np.log(total_n) + 2 * best_nll

    result_dict = {
        'participant_id': participant_id,
        'best_nll': best_nll,
        'best_initial_guess': best_initial_guess,
        'best_parameters': best_parameters,
        'best_EV': best_EV,
        'best_RT': best_RT,
        'AIC': aic,
        'BIC': bic
    }

    print(f"Participant {participant_id} fitted in {(time.time() - start_time) / 60} minutes.")

    return result_dict


class VisualSearchModels:
    def __init__(self, model_type, task='VS', num_params=2, num_options=2):
        """
        Initialize the Model.

        Parameters:
        - reward_means: List of mean rewards for each option.
        - reward_sd: List of standard deviations for each option.
        - model_type: Type of the model.
        - condition: Condition of the model.
        - num_trials: Number of trials for the simulation.
        - num_params: Number of parameters for the model.
        """

        self.model_initialization = None
        self.skip_first = None
        self.num_options = 2 # This is only a placeholder, it will be set in the fit function
        self.initial_EV = None
        self.initial_RT = None
        self.num_training_trials = None
        self.num_exp_restart = None
        self.num_params = num_params
        self.choices_count = np.zeros(self.num_options)
        self.possible_options = [0, 1]
        self.memory_weights = []
        self.choice_history = []
        self.reward_history = []
        self.chosen_history = {option: [] for option in self.possible_options}
        self.reward_history_by_option = {option: [] for option in self.possible_options}
        self.AllProbs = []
        self.PE = []
        self.iteration = 0
        self.epsilon = 1e-12

        # define for each model_type a dict of { attr_name: param_index }
        self._PARAM_MAP = {
            'delta': {'t': 0, 'a': 1},
            'delta_perseveration': {'t': 0, 'a': 1, 'b': 2},
            'delta_asymmetric': {'t': 0, 'a': 1, 'b': 2},
            'delta_RPUT': {'t': 0, 'a': 1},
            'delta_RPUT_unc': {'t': 0, 'a': 1, 'lamda': 2},
            'delta_PVL': {'t': 0, 'a': 1, 'b': 2, 'lamda': 3},
            'delta_PVL_relative': {'t': 0, 'a': 1, 'b': 2, 'lamda': 3},
            'delta_uncertainty': {'t': 0, 'a': 1, 'lamda': 2},
            'decay': {'t': 0, 'a': 1},
            'decay_fre': {'t': 0, 'a': 1, 'b': 2},
            'decay_PVL_relative': {'t': 0, 'a': 1, 'b': 2, 'lamda': 3},
            'decay_choice': {'t': 0, 'a': 1},
            'decay_win': {'t': 0, 'a': 1},
            'decay_RPUT': {'t': 0, 'a': 1},
            'decay_RPUT_unc': {'t': 0, 'a': 1, 'lamda': 2},
            'decay_uncertainty': {'t': 0, 'a': 1, 'lamda': 2},
            'delta_decay': {'t': 0, 'a': 1, 'b': 2},
            'mean_var': {'t': 0, 'lamda': 1},
            'mean_var_delta': {'t': 0, 'a': 1, 'lamda': 2},
            'mean_var_decay': {'t': 0, 'a': 1, 'lamda': 2},
            'mean_var_unc': {'t': 0, 'a': 1, 'lamda': 2},
            'kalman_filter': {'t': 0, 'var_initial': 1},
            'sampler_decay': {'t': 0, 'a': 1},
            'sampler_decay_PE': {'t': 0, 'a': 1},
            'sampler_decay_AV': {'t': 0, 'a': 1},
            'ACTR_Ori': { 'a': 1, 's': 0, 'tau': 2},
            'ACTR': {'t': 0, 'a': 1, 'tau': 2},
            'perseveration': {'w': 0},
            'WSLS': {'p_ws': 0, 'p_ls': 1},
            'WSLS_delta': {'a': 1, 'p_ws': 0, 'p_ls': 2},
            'WSLS_decay_weight': {'t': 0, 'a': 1, 'p_ws': 2, 'p_ls': 3, 'w': 4},
            'WSLS_delta_weight': {'t': 0, 'a': 1, 'p_ws': 2, 'p_ls': 3, 'w': 4},
            'RT_exp_basic': {'t': 0, 'k': 1, 'RT_initial': 2},
            'RT_delta': {'t': 0, 'a': 1, 'RT_initial': 2},
            'RT_decay': {'t': 0, 'a': 1, 'RT_initial': 2},
            'RT_exp_delta': {'t': 0, 'a': 1, 'RT_initial': 2, 'k': 3},
            'RT_delta_PVL': {'t': 0, 'a': 1, 'RT_initial': 2, 'b': 3, 'lamda': 4},
            'hybrid_delta_delta': {'t': 0, 'a': 1, 'RT_initial': 2, 'w': 3},
            'hybrid_delta_delta_3': {'t': 0, 'a': 1, 'b': 2, 'RT_initial': 3, 'w': 4},
            'hybrid_decay_delta': {'t': 0, 'a': 1, 'RT_initial': 2, 'w': 3},
            'hybrid_decay_delta_3': {'t': 0, 'a': 1, 'b': 2, 'RT_initial': 3, 'w': 4},
            'hybrid_decay_decay': {'t': 0, 'a': 1, 'RT_initial': 2, 'w': 3},
            'hybrid_decay_decay_3': {'t': 0, 'a': 1, 'b': 2, 'RT_initial': 3, 'w': 4},
            'hybrid_WSLS_delta': {'t': 0, 'a': 1, 'p_ws': 2, 'p_ls': 3, 'RT_initial': 4, 'w': 5},
        }

        # any attributes you always want to have, even if None
        self._DEFAULT_ATTRS = [
            'RT_initial_suboptimal', 'RT_initial_optimal', 'RT_initial',
            'k', 's', 't', 'a', 'b', 'tau', 'lamda',
            'p_ws', 'p_ls', 'w', 'var_initial', 'unc'
        ]

        # initialize default b overrides
        self._B_OVERRIDES = {
            'hybrid_delta_delta': 'a',
            'hybrid_decay_delta': 'a',
            'hybrid_WSLS_delta': 'a',
            'hybrid_decay_decay': 'a',
        }

        # initialize all attributes to None
        for attr in self._DEFAULT_ATTRS:
            setattr(self, attr, None)

        self.EVs = None
        self.RTs = None
        self.Probs = np.full(self.num_options, 0.25)
        self.mean = None
        self.var = np.full(self.num_options, 1 / 12)
        self.AV = None
        self.RT_AV = None

        # Model type
        self.model_type = model_type
        self.task = task

        # Mapping of updating functions to model types
        self.updating_mapping = {
            'delta': self.delta_update,
            'delta_perseveration': self.delta_perseveration_update,
            'delta_asymmetric': self.delta_asymmetric_update,
            'delta_RPUT': self.delta_reward_per_RT_update,
            'delta_RPUT_unc': self.delta_reward_RPUT_unc_update,
            'delta_PVL': self.delta_PVL_update,
            'delta_PVL_relative': self.delta_PVL_relative_update,
            'delta_uncertainty': self.delta_uncertainty_update,
            'decay': self.decay_update,
            'decay_fre': self.decay_fre_update,
            'decay_PVL_relative': self.decay_PVL_relative_update,
            'decay_choice': self.decay_choice_update,
            'decay_win': self.decay_win_update,
            'decay_RPUT': self.decay_reward_per_RT_update,
            'decay_RPUT_unc': self.decay_reward_RPUT_unc_update,
            'decay_uncertainty': self.decay_uncertainty_update,
            'delta_decay': self.delta_update,
            'mean_var': self.mean_var_update,
            'mean_var_delta': self.mean_var_delta_update,
            'mean_var_decay': self.mean_var_decay_update,
            'mean_var_unc': self.mean_var_decay_unc_update,
            'kalman_filter': self.kalman_filter_update,
            'sampler_decay': self.sampler_decay_update,
            'sampler_decay_PE': self.sampler_decay_PE_update,
            'sampler_decay_AV': self.sampler_decay_AV_update,
            'perseveration': self.perseveration_update,
            'WSLS': self.WSLS_update,
            'WSLS_delta': self.WSLS_delta_update,
            'WSLS_delta_weight': self.WSLS_delta_weight_update,
            'WSLS_decay_weight': self.WSLS_decay_weight_update,
            'RT_exp_basic': self.RT_exp_basic,
            'RT_delta': self.RT_delta,
            'RT_decay': self.RT_decay,
            'RT_exp_delta': self.RT_exp_delta,
            'RT_delta_PVL': self.RT_delta_PVL,
            'RT_decay_PVL': self.RT_decay_PVL,
            'hybrid_delta_delta': self.hybrid_delta_delta,
            'hybrid_delta_delta_3': self.hybrid_delta_delta,
            'hybrid_decay_delta': self.hybrid_decay_delta,
            'hybrid_decay_delta_3': self.hybrid_decay_delta,
            'hybrid_WSLS_delta': self.hybrid_WSLS_delta,
            'hybrid_decay_decay': self.hybrid_decay_decay,
            'hybrid_decay_decay_3': self.hybrid_decay_decay,
        }

        self.updating_function = self.updating_mapping[self.model_type]

        # Mapping of nll functions to model types
        self.nll_mapping_VS = {
            'delta': self.standard_nll,
            'delta_perseveration': self.standard_nll,
            'delta_asymmetric': self.standard_nll,
            'delta_RPUT': self.standard_nll,
            'delta_RPUT_unc': self.standard_nll,
            'delta_PVL': self.standard_nll,
            'delta_PVL_relative': self.standard_nll,
            'delta_uncertainty': self.standard_nll,
            'decay': self.standard_nll,
            'decay_fre': self.standard_nll,
            'decay_PVL_relative': self.standard_nll,
            'decay_choice': self.standard_nll,
            'decay_win': self.standard_nll,
            'decay_RPUT': self.standard_nll,
            'decay_RPUT_unc': self.standard_nll,
            'decay_uncertainty': self.standard_nll,
            'delta_decay': self.standard_nll,
            'mean_var': self.standard_nll,
            'mean_var_delta': self.standard_nll,
            'mean_var_decay': self.standard_nll,
            'mean_var_unc': self.standard_nll,
            'kalman_filter': self.standard_nll,
            'sampler_decay': self.standard_nll,
            'sampler_decay_PE': self.standard_nll,
            'sampler_decay_AV': self.standard_nll,
            'perseveration': self.WSLS_nll,
            'WSLS': self.WSLS_nll,
            'WSLS_delta': self.WSLS_nll,
            'WSLS_delta_weight': self.WSLS_nll,
            'WSLS_decay_weight': self.weight_nll,
            'RT_exp_basic': self.standard_nll,
            'RT_delta': self.standard_nll,
            'RT_decay': self.standard_nll,
            'RT_exp_delta': self.standard_nll,
            'RT_delta_PVL': self.standard_nll,
            'RT_decay_PVL': self.standard_nll,
            'hybrid_delta_delta': self.hybrid_nll,
            'hybrid_delta_delta_3': self.hybrid_nll,
            'hybrid_decay_delta': self.hybrid_nll,
            'hybrid_decay_delta_3': self.hybrid_nll,
            'hybrid_WSLS_delta': self.hybrid_WSLS_nll,
            'hybrid_decay_decay': self.hybrid_nll,
            'hybrid_decay_decay_3': self.hybrid_nll,
        }

        self.nll_function = self.nll_mapping_VS[self.model_type]

    def reset(self):
        """
        Reset the model and empty all the lists.
        """
        self.choices_count = np.zeros(self.num_options)
        self.memory_weights = []
        self.choice_history = []
        self.reward_history = []
        self.chosen_history = {option: [] for option in self.possible_options}
        self.reward_history_by_option = {option: [] for option in self.possible_options}
        self.AllProbs = []
        self.PE = []

        self.EVs = self.initial_EV.copy()
        self.RTs = self.initial_RT.copy()
        self.Probs = np.full(self.num_options, 0.25)
        self.AV = np.mean(self.initial_EV)
        self.RT_AV = np.mean(self.initial_RT)
        self.mean = self.initial_EV.copy()
        self.var = np.full(self.num_options, self.var_initial if self.var_initial is not None else 1 / 12)

    def softmax(self, x):
        c = 3 ** self.t - 1
        e_x = np.exp(np.clip(c * x, -700, 700))
        return np.clip(e_x / e_x.sum(), 1e-12, 1 - 1e-12)

    def simulation_unpacker(self, dict):
        all_sims = []

        for res in dict:
            sim_num = res['simulation_num']
            a_val = res['a']
            b_val = res['b']
            t_val = res['t']
            lambda_val = res['lamda']
            tau_val = res['tau']
            for trial_idx, trial_detail in zip(res['trial_indices'], res['trial_details']):
                data_row = {
                    'simulation_num': sim_num,
                    'trial_index': trial_idx,
                    'a': a_val,
                    'b': b_val,
                    't': t_val,
                    'tau': tau_val,
                    'lambda': lambda_val,
                    'pair': trial_detail['pair'],
                    'choice': trial_detail['choice'],
                    'reward': trial_detail['reward'],
                }
                all_sims.append(data_row)

        return pd.DataFrame(all_sims)

    # ==================================================================================================================
    # Now we define the updating function for each model
    # ==================================================================================================================
    def delta_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.EVs[chosen]
        self.EVs[chosen] += self.a * prediction_error

    def delta_asymmetric_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.EVs[chosen]
        self.EVs[chosen] += (prediction_error > 0) * self.a * prediction_error + (prediction_error < 0) * self.b * \
                            prediction_error

    def delta_perseveration_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.EVs[chosen]
        self.EVs[chosen] += self.a * prediction_error + self.b

    def delta_reward_per_RT_update(self, chosen, reward, rt, trial):
        reward_per_RT = reward / np.clip(rt, 0.0001, None)  # Avoid division by zero
        prediction_error = reward_per_RT - self.EVs[chosen]
        self.EVs[chosen] += self.a * prediction_error

    def delta_reward_RPUT_unc_update(self, chosen, reward, rt, trial):
        reward_per_RT = reward / np.clip(rt, 0.0001, None)  # Avoid division by zero
        prediction_error = reward_per_RT - self.mean[chosen]
        self.mean[chosen] += self.a * prediction_error
        self.var[chosen] += self.a * (prediction_error ** 2 - self.var[chosen])
        self.EVs[chosen] = self.mean[chosen] - (self.lamda * self.var[chosen]) / 2

    def delta_PVL_update(self, chosen, reward, rt, trial):
        utility = (np.abs(reward) ** self.b) * (reward >= 0) + (-self.lamda * (np.abs(reward) ** self.b)) * (reward < 0)
        prediction_error = utility - self.EVs[chosen]
        self.EVs[chosen] += self.a * prediction_error

    def delta_PVL_relative_update(self, chosen, reward, rt, trial):
        reward_diff = reward - self.AV
        utility = (np.abs(reward_diff) ** self.b) * (reward_diff >= 0) + (-self.lamda * (np.abs(reward_diff) ** self.b)) * (reward_diff < 0)
        prediction_error = utility - self.EVs[chosen]
        self.EVs[chosen] += self.a * prediction_error
        self.AV += self.a * reward_diff

    def delta_uncertainty_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.EVs[chosen]
        self.var += self.a * (prediction_error ** 2 - self.var)
        self.EVs[chosen] += self.a * (prediction_error - self.lamda * self.var[chosen])

    def decay_update(self, chosen, reward, rt, trial):
        self.EVs = self.EVs * (1 - self.a)
        self.EVs[chosen] += reward

    def decay_fre_update(self, chosen, reward, rt, trial):
        self.EVs = self.EVs * (1 - self.a)
        self.choices_count[chosen] += 1
        multiplier = self.choices_count[chosen] ** (-self.b)
        self.EVs[chosen] += reward * multiplier

    def decay_choice_update(self, chosen, reward, rt, trial):
        self.EVs = self.EVs * (1 - self.a)
        self.EVs[chosen] += 1

    def decay_PVL_relative_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.AV
        utility = ((np.abs(prediction_error) ** self.b) * (prediction_error >= 0) +
                   (-self.lamda * (np.abs(prediction_error) ** self.b)) * (prediction_error < 0))
        self.EVs = [x * (1 - self.a) for x in self.EVs]
        self.EVs[chosen] += utility
        self.AV += self.a * prediction_error

    def decay_reward_per_RT_update(self, chosen, reward, rt, trial):
        reward_per_RT = reward / np.clip(rt, 0.0001, None)  # Avoid division by zero
        self.EVs = [x * (1 - self.a) for x in self.EVs]
        self.EVs[chosen] += reward_per_RT

    def decay_reward_RPUT_unc_update(self, chosen, reward, rt, trial):
        reward_per_RT = reward / np.clip(rt, 0.0001, None)  # Avoid division by zero
        self.mean = self.mean * (1 - self.a)
        self.var = self.var * (1 - self.a)
        self.choices_count = self.choices_count * (1 - self.a)
        self.choices_count[chosen] += 1
        self.var[chosen] += (reward_per_RT - self.mean[chosen]) ** 2
        uncertainty = [(self.var[x] / np.clip(self.choices_count[x], 1, 9999)) for x in range(self.num_options)]
        self.mean[chosen] += reward_per_RT
        self.EVs = self.mean - (0.5 * self.lamda * np.array(uncertainty))

    def decay_win_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.AV
        self.AV += prediction_error * self.a
        self.EVs = self.EVs * (1 - self.a)
        self.EVs[chosen] += (prediction_error > 0)

    def decay_uncertainty_update(self, chosen, reward, rt, trial):
        self.EVs = self.EVs * (1 - self.a)
        self.var = self.var * (1 - self.a)
        self.var[chosen] += (reward - self.EVs[chosen]) ** 2
        self.EVs[chosen] += reward - self.lamda * self.var[chosen]

    def delta_decay_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.EVs[chosen]
        self.EVs = self.EVs * (1 - self.b)
        self.EVs[chosen] += self.a * prediction_error

    def mean_var_update(self, chosen, reward, rt, trial):
        # no alpha, just average
        prediction_error = reward - self.mean[chosen]
        self.choices_count[chosen] += 1
        self.mean[chosen] += prediction_error / np.clip(self.choices_count[chosen], 1, 9999)
        self.var[chosen] += (prediction_error ** 2 - self.var[chosen]) / np.clip(self.choices_count[chosen], 1, 9999)
        self.EVs[chosen] = self.mean[chosen] - (0.5 * self.lamda * self.var[chosen])

    def mean_var_delta_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.mean[chosen]
        self.mean[chosen] += self.a * prediction_error
        self.var[chosen] += self.a * (prediction_error ** 2 - self.var[chosen])
        self.EVs[chosen] = self.mean[chosen] - (0.5 * self.lamda * self.var[chosen])

    def mean_var_decay_update(self, chosen, reward, rt, trial):
        self.mean = self.mean * (1 - self.a)
        self.var = self.var * (1 - self.a)
        self.var[chosen] += (reward - self.mean[chosen]) ** 2
        self.mean[chosen] += reward
        self.EVs = self.mean - (0.5 * self.lamda * self.var)

    def mean_var_decay_unc_update(self, chosen, reward, rt, trial):
        self.mean = self.mean * (1 - self.a)
        self.var = self.var * (1 - self.a)
        self.choices_count = self.choices_count * (1 - self.a)
        self.choices_count[chosen] += 1
        self.var[chosen] += (reward - self.mean[chosen]) ** 2
        uncertainty = [(self.var[x] / np.clip(self.choices_count[x], 1, 9999)) for x in range(self.num_options)]
        self.mean[chosen] += reward
        self.EVs = self.mean - (0.5 * self.lamda * np.array(uncertainty))

    def kalman_filter_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.EVs[chosen]

        # Kalman Gain
        kalman_gain = self.var[chosen] / (self.var[chosen] + self.var_initial)

        # Update EV and variance
        self.var[chosen] = (1 - kalman_gain) * self.var[chosen]
        self.EVs[chosen] += kalman_gain * prediction_error

    def sampler_decay_update(self, chosen, reward, rt, trial):
        """
        This sampler model is similar to the ACT-R model, but it takes all past trials into account.
        The parameter b consistently decays the weights of past trials over time.
        """

        self.reward_history.append(reward)
        self.choice_history.append(chosen)
        self.memory_weights.append(1)

        # Decay weights of past trials and EVs
        self.EVs = self.EVs * (1 - self.a)
        self.memory_weights = [w * (1 - self.b) for w in self.memory_weights]

        # Compute the probabilities from memory weights
        total_weight = sum(self.memory_weights)
        self.AllProbs = [w / total_weight for w in self.memory_weights]

        # Update EVs based on the samples from memory
        for j in range(len(self.reward_history)):
            self.EVs[self.choice_history[j]] += self.AllProbs[j] * self.reward_history[j]

    def sampler_decay_PE_update(self, chosen, reward, rt, trial):
        """
        In this version of the sampler model, we use prediction errors between the actual reward and the EV of
        the chosen option as reward instead of actual rewards.
        """

        self.reward_history.append(reward)
        self.choice_history.append(chosen)
        self.memory_weights.append(1)

        # use the following code if you want to use prediction errors as reward instead of actual rewards
        prediction_error = self.a * (reward - self.EVs[chosen])
        self.PE.append(prediction_error)

        # Decay weights of past trials and EVs
        self.EVs = self.EVs * (1 - self.a)
        self.memory_weights = [w * (1 - self.b) for w in self.memory_weights]

        # Compute the probabilities from memory weights
        total_weight = sum(self.memory_weights)
        self.AllProbs = [w / total_weight for w in self.memory_weights]

        # Update EVs based on the samples from memory
        for j in range(len(self.reward_history)):
            self.EVs[self.choice_history[j]] += self.AllProbs[j] * self.PE[j]

    def sampler_decay_AV_update(self, chosen, reward, rt, trial):
        """
        In this version of the sampler model, we use prediction errors between the actual reward and the average
        reward over all options as reward instead of actual rewards.
        """

        self.reward_history.append(reward)
        self.choice_history.append(chosen)
        self.memory_weights.append(1)

        # use the following code if you want to use average reward as reward instead of actual rewards
        prediction_error = self.a * (reward - self.AV)
        self.AV += prediction_error
        self.PE.append(prediction_error)

        # Decay weights of past trials and EVs
        self.EVs = self.EVs * (1 - self.a)
        self.memory_weights = [w * (1 - self.b) for w in self.memory_weights]

        # Compute the probabilities from memory weights
        total_weight = sum(self.memory_weights)
        self.AllProbs = [w / total_weight for w in self.memory_weights]

        # Update EVs based on the samples from memory
        for j in range(len(self.reward_history)):
            self.EVs[self.choice_history[j]] += self.AllProbs[j] * self.PE[j]

    def perseveration_update(self, chosen, reward, rt, trial):
        self.Probs[:] = [0.0] * self.num_options
        self.Probs[chosen] = self.w
        n = self.num_options - 1
        for option in range(self.num_options):
            if option != chosen:
                self.Probs[option] = (1 - self.w) / n

    def WSLS_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.AV
        self.AV += prediction_error / (trial + 1)

        pos = int(prediction_error > 0)
        chosen_prob = pos * self.p_ws + (1 - pos) * (1 - self.p_ls)
        other_prob = pos * (1 - self.p_ws) + (1 - pos) * self.p_ls
        self.Probs[:] = [other_prob] * self.num_options
        self.Probs[chosen] = chosen_prob

    def WSLS_delta_update(self, chosen, reward, rt, trial):
        prediction_error = reward - self.AV
        self.AV += prediction_error * self.a

        pos = int(prediction_error > 0)
        chosen_prob = pos * self.p_ws + (1 - pos) * (1 - self.p_ls)
        other_prob = pos * (1 - self.p_ws) + (1 - pos) * self.p_ls
        self.Probs[:] = [other_prob] * self.num_options
        self.Probs[chosen] = chosen_prob

    def WSLS_delta_weight_update(self, chosen, reward, rt, trial):
        # WSLS
        prediction_error = reward - self.AV
        self.AV += prediction_error / (trial + 1)

        pos = int(prediction_error > 0)
        chosen_prob = pos * self.p_ws + (1 - pos) * (1 - self.p_ls)
        other_prob = pos * (1 - self.p_ws) + (1 - pos) * self.p_ls
        self.Probs[:] = [other_prob] * self.num_options
        self.Probs[chosen] = chosen_prob

        # Decay
        prediction_error_per_option = reward - self.EVs[chosen]
        self.EVs[chosen] += self.a * prediction_error_per_option

    def WSLS_decay_weight_update(self, chosen, reward, rt, trial):
        # WSLS
        prediction_error = reward - self.AV
        self.AV += prediction_error / (trial + 1)

        pos = int(prediction_error > 0)
        chosen_prob = pos * self.p_ws + (1 - pos) * (1 - self.p_ls)
        other_prob = pos * (1 - self.p_ws) + (1 - pos) * self.p_ls
        self.Probs[:] = [other_prob] * self.num_options
        self.Probs[chosen] = chosen_prob

        # Decay
        self.EVs = self.EVs * (1 - self.a)
        self.EVs[chosen] += reward

    # ------------------------------------------------------------------------------------------------------------------
    # Now we define RT-related models specifically for the visual search task
    # ------------------------------------------------------------------------------------------------------------------
    def RT_exp_basic(self, chosen, reward, rt, trial):
        self.RTs[chosen] = self.RTs[chosen] * np.exp(-1 * self.k)
        self.EVs = [-x for x in self.RTs]
    
    def RT_delta(self, chosen, reward, rt, trial):
        self.RTs[chosen] += self.a * (rt - self.RTs[chosen])
        self.EVs = [-x for x in self.RTs]

    def RT_delta_PVL(self, chosen, reward, rt, trial):
        RT_diff = rt - self.RT_AV
        self.RT_AV += RT_diff / (trial + 1)
        utility = (self.lamda * np.abs(RT_diff) ** self.b) * (RT_diff > 0) + (-1 * np.abs(RT_diff) ** self.b) * (RT_diff <= 0)
        prediction_error = utility - self.RTs[chosen]
        self.RTs[chosen] += self.a * prediction_error
        self.EVs = [-x for x in self.RTs]

    def RT_decay(self, chosen, reward, rt, trial):
        # force RTs to be negative
        self.RTs = [-np.abs(x) for x in self.RTs]
        self.RTs = [x * (1 - self.a) for x in self.RTs]
        self.RTs[chosen] -= (1 / rt)
        self.EVs = [-x for x in self.RTs]

    def RT_decay_PVL(self, chosen, reward, rt, trial):
        RT_diff = rt - self.RT_AV
        self.RT_AV += RT_diff / (trial + 1)
        utility = (self.lamda * np.abs(rt) ** self.b) * (RT_diff > 0) + (np.abs(rt) ** self.b) * (RT_diff <= 0)
        self.RTs = [-np.abs(x) for x in self.RTs]
        self.RTs = [x * (1 - self.a) for x in self.RTs]
        self.RTs[chosen] -= (1 / utility)
        self.EVs = [-x for x in self.RTs]

    def RT_exp_delta(self, chosen, reward, rt, trial):
        pred = self.RTs[chosen] * np.exp(-1 * self.k)
        prediction_error = rt - pred
        self.RTs[chosen] = pred + self.a * prediction_error
        self.EVs = [-x for x in self.RTs]

    # ------------------------------------------------------------------------------------------------------------------
    # Now we define hybrid models
    # ------------------------------------------------------------------------------------------------------------------
    def hybrid_delta_delta(self, chosen, reward, rt, trial):
        # Reward update
        prediction_error = reward - self.EVs[chosen]
        self.EVs[chosen] += self.a * prediction_error

        # RT update
        self.RTs[chosen] += self.b * (rt - self.RTs[chosen])

    def hybrid_decay_decay(self, chosen, reward, rt, trial):
        # Reward update
        self.EVs = self.EVs * (1 - self.a)
        self.EVs[chosen] += reward

        # RT update
        self.RTs = [-np.abs(x) for x in self.RTs]
        self.RTs = [x * (1 - self.a) for x in self.RTs]
        self.RTs[chosen] -= (1 / rt)

    def hybrid_decay_delta(self, chosen, reward, rt, trial):
        # Reward update
        self.EVs = self.EVs * (1 - self.a)
        self.EVs[chosen] += reward

        # RT update
        self.RTs[chosen] += self.b * (rt - self.RTs[chosen])

    def hybrid_WSLS_delta(self, chosen, reward, rt, trial):
        # Reward update
        prediction_error = reward - self.AV
        self.AV += prediction_error / (trial + 1)

        pos = int(prediction_error > 0)
        chosen_prob = pos * self.p_ws + (1 - pos) * (1 - self.p_ls)
        other_prob = pos * (1 - self.p_ws) + (1 - pos) * self.p_ls
        self.Probs[:] = [other_prob] * self.num_options
        self.Probs[chosen] = chosen_prob

        # RT update
        self.RTs[chosen] += self.b * (rt - self.RTs[chosen])

    def update(self, chosen, reward, rt, trial):
        """
        Update EVs based on the choice, received reward, and trial number.

        Parameters:
        - chosen: Index of the chosen option.
        - reward: Reward received for the current trial.
        - trial: Current trial number.
        """

        self.updating_function(chosen, reward, rt, trial)

    def simulate(self, reward_means, reward_sd, AB_freq, CD_freq, num_trials=250, num_iterations=1000,
                 beta_lower=-1, beta_upper=1):
        """
        Simulate the EV updates for a given number of trials and specified number of iterations.

        Parameters:
        - num_trials: Number of trials for the simulation.
        - AB_freq: Frequency of appearance for the AB pair in the first 150 trials.
        - CD_freq: Frequency of appearance for the CD pair in the first 150 trials.
        - num_iterations: Number of times to repeat the simulation.
        - beta_lower: Lower bound for the beta parameter.
        - beta_upper: Upper bound for the beta parameter.

        Returns:
        - A list of simulation results.
        """
        all_results = []

        for iteration in range(num_iterations):

            print(f"Iteration {iteration + 1} of {num_iterations}")

            self.reset()

            self.s = np.random.uniform(0.0001, 0.9999) if self.model_type == 'ACTR_Ori' else None
            self.t = np.random.uniform(0, 4.9999)
            self.a = np.random.uniform()  # Randomly set decay parameter between 0 and 1
            self.b = np.random.uniform(beta_lower, beta_upper) if self.num_params == 3 else self.a
            self.tau = np.random.uniform(-1.9999, -0.0001) if self.model_type in ('ACTR', 'ACTR_Ori') else None
            self.lamda = np.random.uniform(0.0001, 0.9999) if self.model_type == 'mean_var_delta' else None
            self.choices_count = np.zeros(self.num_options)

            EV_history = np.zeros((num_trials, self.num_options))
            trial_details = []
            trial_indices = []

            training_trials = [(0, 1), (2, 3)]
            training_trial_sequence = [training_trials[0]] * AB_freq + [training_trials[1]] * CD_freq
            np.random.shuffle(training_trial_sequence)

            # Distributing the next 100 trials equally among the four pairs (AC, AD, BC, BD)
            transfer_trials = [(2, 0), (1, 3), (0, 3), (2, 1)]
            transfer_trial_sequence = transfer_trials * 25
            np.random.shuffle(transfer_trial_sequence)

            for trial in range(num_trials):
                trial_indices.append(trial + 1)

                if trial < 150:
                    pair = training_trial_sequence[trial]
                    optimal, suboptimal = (pair[0], pair[1])
                    prob_optimal = self.softmax_mapping[self.model_type](np.array([self.EVs[optimal],
                                                                                   self.EVs[suboptimal]]))[0]
                    chosen = optimal if np.random.rand() < prob_optimal else suboptimal
                else:
                    pair = transfer_trial_sequence[trial - 150]
                    optimal, suboptimal = (pair[0], pair[1])
                    prob_optimal = self.softmax_mapping[self.model_type](np.array([self.EVs[optimal],
                                                                                   self.EVs[suboptimal]]))[0]
                    chosen = optimal if np.random.rand() < prob_optimal else suboptimal

                reward = np.random.normal(reward_means[chosen], reward_sd[chosen])
                trial_details.append(
                    {"trial": trial + 1, "pair": (chr(65 + pair[0]), chr(65 + pair[1])), "choice": chr(65 + chosen),
                     "reward": reward})

                self.update(chosen, reward, trial + 1, pair)
                EV_history[trial] = self.EVs

            all_results.append({
                "simulation_num": iteration + 1,
                "trial_indices": trial_indices,
                "s": self.s,
                "t": self.t,
                "a": self.a,
                "b": self.b,
                "tau": self.tau,
                "lamda": self.lamda,
                "weight": self.w,
                "trial_details": trial_details,
                "EV_history": EV_history
            })

        return self.simulation_unpacker(all_results)

    # ==================================================================================================================
    # Now we define the negative log likelihood function for the ACT-R model because it requires updating EVs before
    # calculating the likelihood
    # ==================================================================================================================
    def WSLS_nll(self, reward, choice, trial, react_time):

        nll = -np.log(0.5)

        for r, ch, t, rt in zip(reward, choice, trial, react_time):
            prob_choice = self.Probs[ch]
            nll += -np.log(max(self.epsilon, prob_choice))

            # Otherwise, we update the model
            self.update(ch, r, rt, t)

        return nll

    def standard_nll(self, reward, choice, trial, react_time):

        nll = -np.log(0.5)

        for r, ch, t, rt in zip(reward, choice, trial, react_time):
            prob_choice = self.softmax(np.array(self.EVs))[ch]
            nll += -np.log(max(self.epsilon, prob_choice))

            # Otherwise, we update the model
            self.update(ch, r, rt, t)

        return nll

    def weight_nll(self, reward, choice, trial, react_time):

        nll = -np.log(0.5)

        for r, ch, t, rt in zip(reward, choice, trial, react_time):
            prob_choice = self.softmax(np.array(self.EVs))[ch] * self.w + (1 - self.w) * self.Probs[ch]
            nll += -np.log(max(self.epsilon, prob_choice))

            # Otherwise, we update the model
            self.update(ch, r, rt, t)

        return nll

    def hybrid_nll(self, reward, choice, trial, react_time):

        nll = -np.log(0.5)

        for r, ch, t, rt in zip(reward, choice, trial, react_time):
            prob_choice_reward = self.softmax(np.array(self.EVs))[ch]
            RT_EVs = [-rt for rt in self.RTs]
            prob_choice_RT = self.softmax(np.array(RT_EVs))[ch]
            prob_choice = self.w * prob_choice_reward + (1 - self.w) * prob_choice_RT
            nll += -np.log(max(self.epsilon, prob_choice))

            # Otherwise, we update the model
            self.update(ch, r, rt, t)

        return nll

    def hybrid_WSLS_nll(self, reward, choice, trial, react_time):

        nll = -np.log(0.5)

        for r, ch, t, rt in zip(reward, choice, trial, react_time):
            prob_choice_reward = self.Probs[ch]
            RT_EVs = [-rt for rt in self.RTs]
            prob_choice_RT = self.softmax(np.array(RT_EVs))[ch]
            prob_choice = self.w * prob_choice_reward + (1 - self.w) * prob_choice_RT
            nll += -np.log(max(self.epsilon, prob_choice))

            # Otherwise, we update the model
            self.update(ch, r, rt, t)

        return nll

    def negative_log_likelihood(self, params, reward, choice, react_time):
        """
        Calculate the negative log likelihood for the model based on the provided parameters,
        rewards, choices, and reaction times.
        """
        self.reset()

        cfg = self._PARAM_MAP.get(self.model_type, {})
        for attr, idx in cfg.items():
            setattr(self, attr, params[idx])

        # set the initial RTs
        if self.RT_initial is not None:
            if self.model_type in ('RT_decay', 'RT_decay_decay', 'RT_decay_decay_3'):
                self.RTs = np.full(self.num_options, (1 / self.RT_initial))
                self.RT_AV = self.RT_initial
            elif self.model_type in ('RT_decay_PVL'):
                prior_utility = 1 / (np.abs(self.RT_initial) ** self.b)
                self.RTs = np.full(self.num_options, prior_utility)
                self.RT_AV = self.RT_initial
            else:
                self.RTs = np.full(self.num_options, self.RT_initial)
                self.RT_AV = self.RT_initial

        self.b = getattr(self, self._B_OVERRIDES.get(self.model_type, 'b'))

        trial = np.arange(1, len(reward) + 1)

        self.model_initialization(reward, choice, react_time, trial)

        # slice data if necessary
        s = slice(self.skip_first, None)
        reward = reward[s]
        choice = choice[s]
        trial = trial[s]

        return self.nll_function(reward, choice, trial, react_time)

    def fixed_init(self, reward, choice, react_time, trial):
        pass

    def first_trial_init(self, reward, choice, react_time, trial):
        # Initialize the model on the first trial
        self.update(choice[0], reward[0], react_time[0], trial[0])

        # Populate the EVs for the first trial
        EV_trial1 = self.EVs[choice[0]]
        RT_trial1 = self.RTs[choice[0]]
        self.EVs = np.full(self.num_options, EV_trial1)
        self.RTs = np.full(self.num_options, RT_trial1)
        self.AV = EV_trial1
        self.RT_AV = RT_trial1

    def first_trial_no_alpha_init(self, reward, choice, react_time, trial):
        # Temporarily force alpha=1 so that update adds full PE
        orig_a = self.a
        self.a = 1.0

        # Initialize the model on the first trial
        self.update(choice[0], reward[0], react_time[0], trial[0])

        # Restore the original alpha value
        self.a = orig_a

        # Populate the EVs for the first trial
        self.EVs = np.full(self.num_options, self.EVs[choice[0]])
        self.RTs = np.full(self.num_options, self.RTs[choice[0]])
        self.var = np.full(self.num_options, self.var[choice[0]])
        self.mean = np.full(self.num_options, self.mean[choice[0]])
        self.AV = self.EVs[choice[0]]
        self.RT_AV = self.RTs[choice[0]]
        self.Probs = np.full(self.num_options, 0.25)

    def fit(self, data, num_training_trials=999, num_exp_restart=999, num_iterations=100, initial_EV=None,
            initial_RT=None, initial_mode='fixed', beta_lower=-1, beta_upper=1):

        self.num_training_trials = num_training_trials
        self.num_exp_restart = num_exp_restart
        self.num_options = len(initial_EV) if initial_EV is not None else 2
        self.initial_EV = np.array(initial_EV or [0.0, 0.0], dtype=float)
        self.initial_RT = np.array(initial_RT or [0.0, 0.0], dtype=float)

        if initial_mode == 'fixed':
            self.model_initialization = self.fixed_init
        elif initial_mode == 'first_trial':
            self.model_initialization = self.first_trial_init
        elif initial_mode == 'first_trial_no_alpha':
            self.model_initialization = self.first_trial_no_alpha_init
        self.skip_first = 0 if initial_mode == 'fixed' else 1

        # Creating a list to hold the future results
        futures = []
        results = []

        # Starting a pool of workers with ProcessPoolExecutor
        with ProcessPoolExecutor(max_workers=16) as executor:
            # Submitting jobs to the executor for each participant
            for participant_id, participant_data in data.items():
                # fit_participant is the function to be executed in parallel
                future = executor.submit(fit_participant, self, participant_id, participant_data, self.model_type,
                                         num_iterations, beta_lower, beta_upper)
                futures.append(future)

            # Collecting results as they complete
            for future in futures:
                results.append(future.result())

        return pd.DataFrame(results)

    def post_hoc_simulation(self, fitting_result, original_data, reward_means,
                            reward_sd, num_iterations=1000, AB_freq=100, CD_freq=50, use_random_sequence=True,
                            sim_trials=250, summary=False):

        num_parameters = len(fitting_result['best_parameters'][0].strip('[]').split())

        parameter_sequences = []
        for i in range(num_parameters):
            sequence = fitting_result['best_parameters'].apply(
                lambda x: float(x.strip('[]').split()[i]) if isinstance(x, str) else np.nan
            )
            parameter_sequences.append(sequence)

        if not use_random_sequence:
            trial_index = original_data.groupby('Subnum')['trial_index'].apply(list)
            trial_sequence = original_data.groupby('Subnum')['TrialType'].apply(list)
        else:
            trial_index, trial_sequence = None, None

        # start the simulation
        all_results = []

        for participant in fitting_result['participant_id']:

            # extract the best-fitting parameters for the current participant
            self.s = parameter_sequences[0][participant - 1] if self.model_type == 'ACTR_Ori' else None
            self.t = parameter_sequences[0][participant - 1]
            self.a = parameter_sequences[1][participant - 1]
            self.b = parameter_sequences[2][participant - 1] if num_parameters == 3 else self.a
            self.tau = parameter_sequences[2][participant - 1] if self.model_type in ('ACTR', 'ACTR_Ori') else None
            self.lamda = parameter_sequences[2][participant - 1] if self.model_type == 'mean_var_delta' else None

            for _ in range(num_iterations):

                print(f"Participant {participant} - Iteration {_ + 1} of {num_iterations}")

                self.reset()

                self.iteration = 0

                trial_details = []
                trial_indices = []

                if use_random_sequence:

                    training_trial_sequence, transfer_trial_sequence = generate_random_trial_sequence(AB_freq, CD_freq)

                    for trial in range(sim_trials):
                        trial_indice = trial + 1
                        trial_indices.append(trial_indice)

                        if trial_indice < 151:
                            pair = training_trial_sequence[trial_indice - 1]
                        else:
                            pair = transfer_trial_sequence[trial_indice - 151]

                        optimal, suboptimal = pair[0], pair[1]

                        prob_optimal = self.softmax_mapping[self.model_type](np.array([self.EVs[optimal],
                                                                                       self.EVs[suboptimal]]))[0]
                        chosen = 1 if np.random.rand() < prob_optimal else 0

                        reward = np.random.normal(reward_means[optimal if chosen == 1 else suboptimal],
                                                  reward_sd[optimal if chosen == 1 else suboptimal])

                        trial_details.append(
                            {"pair": pair, "choice": chosen, "reward": reward}
                        )

                        self.updating_function(optimal if chosen == 1 else suboptimal, reward, trial,
                                               pair)

                else:
                    for trial, pair in zip(trial_index[participant], trial_sequence[participant]):
                        trial_indices.append(trial)

                        optimal, suboptimal = self.choiceset_mapping[1][pair]

                        prob_optimal = self.softmax_mapping[self.model_type](np.array([self.EVs[optimal],
                                                                                       self.EVs[suboptimal]]))[0]
                        chosen = 1 if np.random.rand() < prob_optimal else 0

                        reward = np.random.normal(reward_means[optimal if chosen == 1 else suboptimal],
                                                  reward_sd[optimal if chosen == 1 else suboptimal])

                        trial_details.append(
                            {"pair": pair, "choice": chosen, "reward": reward}
                        )

                        self.update(optimal if chosen == 1 else suboptimal, reward, trial,
                                    self.choiceset_mapping[1][pair])

                all_results.append({
                    "Subnum": participant,
                    "s": self.s,
                    "t": self.t,
                    "a": self.a,
                    "b": self.b,
                    "lamda": self.lamda if self.model_type == 'mean_var_delta' else None,
                    "tau": self.tau if self.model_type in ('ACTR', 'ACTR_Ori') else None,
                    "trial_indices": trial_indices,
                    "trial_details": trial_details
                })

        unpacked_results = []

        for result in all_results:
            sim_num = result['Subnum']

            s = result["s"]
            t = result["t"]
            a = result["a"]
            b = result["b"]
            lamda = result["lamda"] if self.model_type == 'mean_var_delta' else None
            tau = result["tau"] if self.model_type == 'ACTR' else None

            for trial_idx, trial_detail in zip(result['trial_indices'], result['trial_details']):
                var = {
                    "Subnum": sim_num,
                    "trial_index": trial_idx,
                    "s": s,
                    "t": t,
                    "a": a,
                    "b": b,
                    "lamda": lamda,
                    "tau": tau,
                    "pair": trial_detail['pair'],
                    "choice": trial_detail['choice'],
                    "reward": trial_detail['reward'],
                }

                unpacked_results.append(var)

        df = pd.DataFrame(unpacked_results)
        df.dropna(how='all', inplace=True, axis=1)
        if all(df['a'] == df['b']):
            df.drop('b', axis=1, inplace=True)

        if summary:
            if use_random_sequence:
                # remove the trial index column as it is not needed for the summary
                df.drop('trial_index', axis=1, inplace=True)
                summary_df = df.groupby(['Subnum', 'pair']).mean().reset_index()
            else:
                summary_df = df.groupby(['Subnum', 'trial_index']).mean().reset_index()

            return summary_df

        else:
            return df

    def bootstrapping_post_hoc_simulation(self, fitting_result, reward_means, reward_sd, num_iterations=1000,
                                          AB_freq=100, CD_freq=50, sim_trials=250, summary=False):

        num_parameters = len(fitting_result['best_parameters'][0].strip('[]').split())

        parameter_sequences = []
        for i in range(num_parameters):
            sequence = fitting_result['best_parameters'].apply(
                lambda x: float(x.strip('[]').split()[i]) if isinstance(x, str) else np.nan
            )
            parameter_sequences.append(sequence)

        # start the simulation
        all_results = []

        for n in range(num_iterations):
            print(f"Iteration {n + 1} of {num_iterations}")

            # randomly sample with replacement from the original data
            random_idx = random.randint(0, len(fitting_result) - 1)
            self.t = parameter_sequences[0][random_idx]

            self.s = parameter_sequences[0][random_idx] if self.model_type == 'ACTR_Ori' else None
            self.t = parameter_sequences[0][random_idx]
            self.a = parameter_sequences[1][random_idx]
            self.b = parameter_sequences[2][random_idx] if num_parameters == 3 else self.a
            self.tau = parameter_sequences[2][random_idx] if self.model_type in ('ACTR', 'ACTR_Ori') else None
            self.lamda = parameter_sequences[2][random_idx] if self.model_type == 'mean_var_delta' else None

            self.reset()

            trial_details = []
            trial_indices = []

            training_trial_sequence, transfer_trial_sequence = generate_random_trial_sequence(AB_freq, CD_freq)

            for trial in range(sim_trials):
                trial_indice = trial + 1
                trial_indices.append(trial_indice)

                if trial_indice < 151:
                    pair = training_trial_sequence[trial_indice - 1]
                else:
                    pair = transfer_trial_sequence[trial_indice - 151]

                optimal, suboptimal = pair[0], pair[1]

                prob_optimal = self.softmax_mapping[self.model_type](np.array([self.EVs[optimal],
                                                                               self.EVs[suboptimal]]))[0]
                chosen = 1 if np.random.rand() < prob_optimal else 0

                reward = np.random.normal(reward_means[optimal if chosen == 1 else suboptimal],
                                          reward_sd[optimal if chosen == 1 else suboptimal])

                trial_details.append(
                    {"pair": pair, "choice": chosen, "reward": reward}
                )

                self.updating_function(optimal if chosen == 1 else suboptimal, reward, trial,
                                       pair)

            all_results.append({
                "Subnum": n + 1,
                "s": self.s,
                "t": self.t,
                "a": self.a,
                "b": self.b,
                "lamda": self.lamda if self.model_type == 'mean_var_delta' else None,
                "tau": self.tau if self.model_type in ('ACTR', 'ACTR_Ori') else None,
                "trial_indices": trial_indices,
                "trial_details": trial_details
            })

        unpacked_results = []

        for result in all_results:
            subnum = result['Subnum']
            s = result["s"]
            t = result["t"]
            a = result["a"]
            b = result["b"]
            lamda = result["lamda"] if self.model_type == 'mean_var_delta' else None
            tau = result["tau"] if self.model_type == 'ACTR' else None

            for trial_idx, trial_detail in zip(result['trial_indices'], result['trial_details']):
                var = {
                    "Subnum": subnum,
                    "trial_index": trial_idx,
                    "s": s,
                    "t": t,
                    "a": a,
                    "b": b,
                    "lamda": lamda,
                    "tau": tau,
                    "pair": trial_detail['pair'],
                    "choice": trial_detail['choice'],
                    "reward": trial_detail['reward'],
                }

                unpacked_results.append(var)

        df = pd.DataFrame(unpacked_results)
        df.dropna(how='all', inplace=True, axis=1)
        if all(df['a'] == df['b']):
            df.drop('b', axis=1, inplace=True)

        if summary:
            df.drop('trial_index', axis=1, inplace=True)
            summary_df = df.groupby(['Subnum', 'pair']).mean().reset_index()
            return summary_df
        else:
            return df

# ======================================================================================================================
# Assistant functions for the VisualSearchModels class
# ======================================================================================================================
def behavioral_moving_window(data, variable='Optimal_Choice', window_size=10, exclusionary_criteria=None):
    mv_results = []

    for _, participant_data in data.groupby('SubNo'):
        for i in range(len(participant_data) - window_size + 1):
            window = participant_data.iloc[i:i + window_size]
            window = exclusionary_criteria(window)
            optimal_percent = np.mean(window[variable])
            mv_results.append({
                'participant_id': participant_data['SubNo'].iloc[0],
                'window_id': i + 1,
                'optimal_percentage': optimal_percent,
                'Group': participant_data['Group'].iloc[0]
            })

    optimal_window_df = pd.DataFrame(mv_results)
    optimal_window_df['Group'] = optimal_window_df['Group'].map(
        {1: 'High-Reward-Optimal', 2: 'Low-Reward-Optimal'})
    return optimal_window_df

def create_model_summary_df(model_results, criteria='BIC', return_best=False):
    # Find the best model for each participant and group based on BIC
    all_bics = pd.concat((df[['participant_id','Group',criteria]].assign(Model=name)
                          for name, df in model_results.items()), ignore_index=True)
    all_bics = all_bics.dropna(subset=['BIC']) # Ensure BIC is not NaN
    best_idx = all_bics.groupby(['participant_id','Group'])[criteria].idxmin()
    best = all_bics.loc[best_idx]

    # Count how many times each (Model, Group) is best
    best_counts = (best.groupby(['Model','Group']).size().rename('N_Best_Fit').reset_index())

    # Calculate average AIC and BIC for each model and group
    summary_stats = pd.concat((df.groupby('Group').agg(
        Group=('Group', 'first'),
        AIC=('AIC', 'mean'),
        BIC=('BIC', 'mean')).assign(Model=name) for name, df in model_results.items()), ignore_index=True)

    # Merge summary statistics with best counts
    result = (summary_stats.merge(best_counts, on=['Model','Group'], how='left').fillna({'N_Best_Fit': 0}).sort_values(['Model','Group']).reset_index(drop=True))

    if return_best:
        return result, best
    else:
        return result


def create_model_summary_table(model_results, output_path,
                               group_names=['High-Reward-Optimal', 'Low-Reward-Optimal']):
    # Create model summary dataframe
    model_summary_df = create_model_summary_df(model_results)

    # Create summary table in DOCX format
    doc = Document()
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Group'
    header_cells[1].text = 'Model'
    header_cells[2].text = 'AIC'
    header_cells[3].text = 'BIC'
    header_cells[4].text = 'N Best Fit'

    # Find best models for each group based on BIC
    best_models = {}
    n_group = model_summary_df['Group'].nunique()
    for group in range(1, n_group + 1):
        group_data = model_summary_df[model_summary_df['Group'] == group]
        if group_data.empty:
            continue
        best_model_row = group_data.loc[group_data['BIC'].idxmin()]
        best_models[group] = best_model_row

    # Add data rows for each group
    for group in range(1, n_group + 1):
        group_data_all = model_summary_df[model_summary_df['Group'] == group]
        if group_data_all.empty:
            continue

        # Group header row
        row_cells = table.add_row().cells
        # Use provided group name if available, otherwise fallback to "Group X"
        if group - 1 < len(group_names):
            group_label = group_names[group - 1]
        else:
            group_label = f'Group {group}'
        row_cells[0].text = group_label
        row_cells[1].text = ''
        row_cells[2].text = ''
        row_cells[3].text = ''
        row_cells[4].text = ''

        # Rows per model within this group
        for model in model_results.keys():
            group_data = model_summary_df[
                (model_summary_df['Model'] == model) &
                (model_summary_df['Group'] == group)
            ]

            if group_data.empty:
                continue

            row_cells = table.add_row().cells
            row_cells[0].text = ''
            row_cells[1].text = model
            row_cells[2].text = f"{group_data['AIC'].values[0]:.2f}"
            row_cells[3].text = f"{group_data['BIC'].values[0]:.2f}"
            row_cells[4].text = f"{int(group_data['N_Best_Fit'].values[0])}"

            # Highlight best model row for this group
            if group in best_models and model == best_models[group]['Model']:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Save the document
    doc.save(output_path)
# ======================================================================================================================
# End of the VisualSearchModels class (Other assistant classes can be found in ComputationalModeling.py)
# ======================================================================================================================